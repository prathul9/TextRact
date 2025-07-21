
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from ExtractTable import ExtractTable
from docx import Document
import google.generativeai as genai
from PIL import Image
import io
import base64
import os
import sqlite3
from datetime import datetime
from pdf2image import convert_from_path

app = Flask(_name_)
CORS(app, supports_credentials=True)

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "output"
DB_PATH = "upload_history.db"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


et_sess = ExtractTable(api_key="zt7Sf5wdRBiYqaKlEXaV9rH1BR40uJUPsPFtVkrk")


genai.configure(api_key="AIzaSyDs5tpMIcrqQyaOIeCa5RBZl1NgLoWrDgM")
model = genai.GenerativeModel("gemini-1.5-flash")


def init_db():
    with sqlite3.connect(DB_PATH) as conn:
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS uploads (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT NOT NULL,
                filename TEXT NOT NULL,
                timestamp TEXT NOT NULL
            )
        """)
        conn.commit()

init_db()

@app.route("/upload", methods=["POST"])
def upload_image():
    if "image" not in request.files:
        return jsonify({"error": "No image uploaded"}), 400

    username = request.form.get("username")
    if not username:
        return jsonify({"error": "Username is required"}), 400

    file = request.files["image"]
    if file.filename == "":
        return jsonify({"error": "Empty filename"}), 400

    file_ext = file.filename.split(".")[-1].lower()
    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(file_path)

    images = []

    try:
        if file_ext == "pdf":
            images = convert_from_path(file_path, dpi=300)
        else:
            images.append(Image.open(file_path))
    except Exception as e:
        return jsonify({"error": f"Failed to open file: {str(e)}"}), 400

    extracted_texts = []

    for img in images:
        try:
            img_bytes = io.BytesIO()
            img.save(img_bytes, format="PNG")
            image_b64 = base64.b64encode(img_bytes.getvalue()).decode("utf-8")

            gemini_response = model.generate_content([
                {"text": "Extract only non-tabular text from this image."},
                {"mime_type": "image/png", "data": image_b64}
            ])

            extracted_text = gemini_response.text if gemini_response and gemini_response.text else "No text found."
        except Exception as e:
            extracted_text = f"Error extracting text: {str(e)}"

        extracted_texts.append(extracted_text)

    table_data_list = []
    try:
        table_data_list = et_sess.process_file(filepath=file_path, output_format="df") or []
    except Exception as e:
        print("ExtractTable Error:", e)

    doc = Document()
    doc.add_heading("Extracted Text", level=1)
    for text in extracted_texts:
        doc.add_paragraph(text)

    if table_data_list:
        doc.add_heading("Extracted Tables", level=1)
        for index, df in enumerate(table_data_list):
            doc.add_paragraph(f"Table {index + 1}")
            table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
            table.style = "Table Grid"

            hdr_cells = table.rows[0].cells
            for j, col_name in enumerate(df.columns):
                hdr_cells[j].text = str(col_name)

            for i, row in df.iterrows():
                row_cells = table.rows[i + 1].cells
                for j, value in enumerate(row):
                    row_cells[j].text = str(value)

    doc_filename = f"{os.path.splitext(file.filename)[0]}_output.docx"
    doc_path = os.path.join(OUTPUT_FOLDER, doc_filename)
    doc.save(doc_path)

    
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with sqlite3.connect(DB_PATH) as conn:
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO uploads (username, filename, timestamp) VALUES (?, ?, ?)",
            (username, doc_filename, timestamp)
        )
        conn.commit()

    return jsonify({
        "text": "\n".join(extracted_texts),
        "doc_url": f"http://127.0.0.1:5000/download/{doc_filename}"
    }), 200

@app.route("/history/<username>", methods=["GET"])
def get_history(username):
    with sqlite3.connect(DB_PATH) as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT filename, timestamp FROM uploads WHERE username = ? ORDER BY id DESC", (username,))
        rows = cursor.fetchall()

    history = [{"filename": r[0], "timestamp": r[1]} for r in rows]
    return jsonify({"history": history})

@app.route("/download/<filename>", methods=["GET"])
def download_file(filename):
    filepath = os.path.join(OUTPUT_FOLDER, filename)
    if not os.path.exists(filepath):
        return jsonify({"error": "File not found"}), 404
    return send_file(filepath, as_attachment=True)

if _name_ == "_main_":
    app.run(debug=True)
